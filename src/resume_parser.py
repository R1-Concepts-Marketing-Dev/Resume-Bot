"""Extract plain text from a resume attachment. Tries fast text extraction
first (pypdf for PDF, python-docx for Word, plain decode for TXT/RTF, Claude
vision for images). If a PDF returns very little text (likely a scanned
image), falls back to OCR via Tesseract."""

from __future__ import annotations

import base64
import io
import logging
import re

log = logging.getLogger(__name__)

# If pypdf extracts fewer than this many characters from a PDF, we treat it
# as image-based and OCR it instead.
OCR_FALLBACK_CHAR_THRESHOLD = 200

# MIME types Anthropic's vision API accepts. HEIC/TIFF/etc. fall through to
# the Tesseract OCR path because Claude can't read them natively yet.
_CLAUDE_VISION_MIME_MAP = {
    "image/jpeg": "image/jpeg",
    "image/jpg":  "image/jpeg",
    "image/png":  "image/png",
    "image/webp": "image/webp",
    "image/gif":  "image/gif",
}

# Cheap, fast model for OCR-style image-to-text extraction.
_VISION_MODEL = "claude-haiku-4-5"

# Prompt for the vision extraction call. Kept deliberately minimal -- we
# want raw text, not commentary or formatting.
_VISION_PROMPT = (
    "This is a resume image. Extract ALL text from it verbatim, preserving "
    "section order and bullets where you can. Do not summarize, paraphrase, "
    "or add any commentary -- just return the raw text content. If the image "
    "is not a resume, return the text you do see plus a single line at the "
    "top saying: [NOT A RESUME]."
)


def _extract_pdf_text(data: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    chunks = []
    for page in reader.pages:
        try:
            chunks.append(page.extract_text() or "")
        except Exception as e:
            log.warning("pypdf page extract failed: %s", e)
    return "\n".join(chunks).strip()


def _ocr_pdf(data: bytes) -> str:
    """Rasterize each page to an image and OCR. Requires tesseract + poppler
    installed on the runner."""
    try:
        from pdf2image import convert_from_bytes
        import pytesseract
    except ImportError as e:
        log.warning("OCR dependencies unavailable: %s", e)
        return ""
    try:
        pages = convert_from_bytes(data, dpi=200)
    except Exception as e:
        log.warning("pdf2image failed: %s", e)
        return ""
    out = []
    for img in pages:
        try:
            out.append(pytesseract.image_to_string(img) or "")
        except Exception as e:
            log.warning("OCR page failed: %s", e)
    return "\n".join(out).strip()


def _extract_docx_text(data: bytes) -> str:
    import docx
    doc = docx.Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text]
    # Also grab table cell text -- resumes often use tables for layout
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts).strip()


def _extract_txt(data: bytes) -> str:
    """Plain text or RTF-lite attachment. Tries utf-8 first, then utf-16
    and latin-1 as fallbacks for older Windows-saved files."""
    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            text = data.decode(enc, errors="strict")
        except (UnicodeDecodeError, UnicodeError):
            continue
        # Strip RTF control words if it looks like an RTF document
        if text.lstrip().startswith("{\\rtf"):
            text = re.sub(r"\\[a-z]+-?\d* ?", " ", text)
            text = re.sub(r"[{}]", "", text)
        return text.strip()
    # Last resort: lossy decode
    return data.decode("utf-8", errors="replace").strip()


def _ocr_image_tesseract(data: bytes) -> str:
    """Local OCR via Tesseract. Used as a fallback when Claude vision is
    unavailable or doesn't support the image MIME type."""
    try:
        from PIL import Image
        import pytesseract
    except ImportError as e:
        log.warning("Image OCR dependencies unavailable: %s", e)
        return ""
    try:
        img = Image.open(io.BytesIO(data))
        return (pytesseract.image_to_string(img) or "").strip()
    except Exception as e:
        log.warning("Image OCR failed: %s", e)
        return ""


def _extract_image_claude(data: bytes, mime_type: str, api_key: str) -> str:
    """Use Claude's vision API to read a resume image. Returns extracted
    text on success, empty string if the call fails (caller will fall
    back to Tesseract OCR).

    Cost: ~1 cent per image with claude-haiku-4-5 (varies with image size
    and text density). Much better accuracy than Tesseract on phone-camera
    photos, angled scans, and handwriting."""
    if not api_key:
        return ""
    mt = (mime_type or "").lower()
    api_mime = _CLAUDE_VISION_MIME_MAP.get(mt)
    if not api_mime:
        log.info("Image MIME %r not supported by Claude vision; using Tesseract", mime_type)
        return ""

    try:
        import anthropic
    except ImportError as e:
        log.warning("anthropic SDK unavailable, falling back to Tesseract: %s", e)
        return ""

    try:
        client = anthropic.Anthropic(api_key=api_key)
        resp = client.messages.create(
            model=_VISION_MODEL,
            max_tokens=4000,
            messages=[{
                "role": "user",
                "content": [
                    {
                        "type": "image",
                        "source": {
                            "type": "base64",
                            "media_type": api_mime,
                            "data": base64.standard_b64encode(data).decode("ascii"),
                        },
                    },
                    {"type": "text", "text": _VISION_PROMPT},
                ],
            }],
        )
        text = "".join(
            b.text for b in resp.content if getattr(b, "type", "") == "text"
        ).strip()
        return text
    except Exception as e:
        log.warning("Claude vision failed (%s); will fall back to Tesseract", e)
        return ""


def extract(filename: str, mime_type: str, data: bytes,
            *, api_key: str = "") -> tuple[str, bool]:
    """Returns (text, used_ocr). Empty text means we couldn't read it --
    main.py treats that as Unreadable.

    api_key (optional): Anthropic API key. If provided AND the attachment
    is an image type Claude vision supports (JPEG/PNG/WEBP/GIF), uses
    Claude vision instead of Tesseract for better accuracy. Falls back
    to Tesseract on errors or unsupported formats (HEIC, TIFF)."""
    name = (filename or "").lower()
    mt = (mime_type or "").lower()
    is_pdf = "pdf" in mt or name.endswith(".pdf")
    is_docx = "wordprocessingml" in mt or name.endswith(".docx")
    is_txt = (
        mt.startswith("text/plain") or mt.startswith("text/rtf")
        or name.endswith((".txt", ".rtf"))
    )
    is_image = mt.startswith("image/") or name.endswith(
        (".jpg", ".jpeg", ".png", ".heic", ".webp", ".tif", ".tiff")
    )

    if is_pdf:
        text = _extract_pdf_text(data)
        if len(text) >= OCR_FALLBACK_CHAR_THRESHOLD:
            return text, False
        log.info("PDF text extraction returned %d chars; falling back to OCR", len(text))
        ocr_text = _ocr_pdf(data)
        return ocr_text or text, bool(ocr_text)

    if is_docx:
        return _extract_docx_text(data), False

    if is_txt:
        return _extract_txt(data), False

    if is_image:
        # Try Claude vision first (better accuracy, especially on photos);
        # fall back to local Tesseract if Claude doesn't support the MIME
        # type, the API call fails, or no api_key was provided.
        claude_text = _extract_image_claude(data, mime_type, api_key)
        if claude_text:
            return claude_text, True
        log.info("Claude vision unavailable for this image; trying Tesseract")
        tesseract_text = _ocr_image_tesseract(data)
        return tesseract_text, bool(tesseract_text)

    # Older .doc (binary Word) isn't handled in pure Python. We could shell
    # out to LibreOffice but that's a 30s startup hit. For now we surface
    # "Unreadable" and let HR pull the resume manually from the email thread.
    log.warning("Unsupported attachment type: %s (%s)", filename, mime_type)
    return "", False
